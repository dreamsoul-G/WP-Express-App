"""
WP快通（文档处理）软件
Copyright [2026] [郭宇轩]

Licensed under the Apache License, Version 2.0 (the "License");
you may not use this file except in compliance with the License.
You may obtain a copy of the License at

    http://www.apache.org/licenses/LICENSE-2.0

Unless required by applicable law or agreed to in writing, software
distributed under the License is distributed on an "AS IS" BASIS,
WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
See the License for the specific language governing permissions and
limitations under the License.
"""

from pathlib import Path
from typing import List, Optional

from .base import BaseHandler

try:
    import docx

    HAVE_DOCX = True
except ImportError:
    HAVE_DOCX = False
    print("警告: python-docx 未安装，Word功能不可用")


class WordHandler(BaseHandler):
    """Word文档处理器"""

    def __init__(self):
        super().__init__()
        self.file_ext = ".docx"

    def create_single_file(self, file_name: str, save_path: Optional[Path] = None) -> bool:
        """创建单个Word文档"""
        if not HAVE_DOCX:
            self.logger.error("python-docx 未安装")
            return False

        try:
            # 确保文件名有扩展名
            if not file_name.endswith(self.file_ext):
                file_name += self.file_ext

            # 确定保存路径
            if save_path:
                output_path = save_path / file_name
            else:
                output_path = self.default_save_path / file_name

            # 确保目录存在
            output_path.parent.mkdir(parents=True, exist_ok=True)

            # 创建文档
            doc = docx.Document()
            doc.add_paragraph("这是一个新Word文档")
            doc.save(output_path)

            self.logger.info(f"创建Word文档: {output_path}")
            return True

        except Exception as e:
            self.logger.error(f"创建Word文档失败: {e}")
            return False

    def create_multiple(self, count: int, prefix: str, save_path: Optional[Path] = None) -> bool:
        """批量创建Word文档"""
        if not HAVE_DOCX:
            self.logger.error("python-docx 未安装")
            return False

        try:
            # 确定保存路径
            if save_path:
                folder_path = save_path / prefix
            else:
                folder_path = self.default_save_path / prefix

            # 创建文件夹
            folder_path.mkdir(parents=True, exist_ok=True)

            # 批量创建
            for i in range(1, count + 1):
                file_name = f"{prefix}_{i}{self.file_ext}"
                file_path = folder_path / file_name

                doc = docx.Document()
                doc.add_paragraph(f"这是第 {i} 个文档")
                doc.save(file_path)

            self.logger.info(f"批量创建 {count} 个Word文档到: {folder_path}")
            return True

        except Exception as e:
            self.logger.error(f"批量创建Word文档失败: {e}")
            return False

    def merge_files(self, source_files: List[Path], output_path: Path) -> bool:
        """合并Word文档"""
        if not HAVE_DOCX:
            self.logger.error("python-docx 未安装")
            return False

        if not source_files:
            self.logger.error("没有源文件")
            return False

        try:
            # 确保输出目录存在
            output_path.parent.mkdir(parents=True, exist_ok=True)

            # 以第一个文档为基础
            base_doc = docx.Document(source_files[0])

            # 合并其他文档
            for file_path in source_files[1:]:
                current_doc = docx.Document(file_path)
                base_doc.add_page_break()
                for element in current_doc.element.body:
                    base_doc.element.body.append(element)

            # 保存合并后的文档
            base_doc.save(output_path)

            self.logger.info(f"合并Word文档到: {output_path}")
            return True

        except Exception as e:
            self.logger.error(f"合并Word文档失败: {e}")
            return False

    def split_file(self, source_path: Path, split_pos: int, output_dir: Optional[Path] = None) -> List[Path]:
        """拆分Word文档"""
        if not HAVE_DOCX:
            self.logger.error("python-docx 未安装")
            return []

        try:
            # 确定输出目录
            if output_dir:
                save_dir = output_dir
            else:
                save_dir = source_path.parent

            # 打开源文档
            source_doc = docx.Document(source_path)
            all_paragraphs = source_doc.element.body

            # 创建拆分文档
            doc1 = docx.Document()
            doc2 = docx.Document()

            # 拆分文档
            current_para = 0
            for element in all_paragraphs:
                if element.tag.endswith('p'):
                    current_para += 1
                    if current_para < split_pos:
                        doc1.element.body.append(element)
                    else:
                        doc2.element.body.append(element)

            # 生成输出路径
            source_name = source_path.stem
            output1 = save_dir / f"{source_name}_拆分1{self.file_ext}"
            output2 = save_dir / f"{source_name}_拆分2{self.file_ext}"

            # 保存文档
            doc1.save(output1)
            doc2.save(output2)

            self.logger.info(f"拆分Word文档: {output1}, {output2}")
            return [output1, output2]

        except Exception as e:
            self.logger.error(f"拆分Word文档失败: {e}")
            return []
